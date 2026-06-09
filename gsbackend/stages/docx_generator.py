"""
Utilitaire de génération de documents Word (.docx)
pour les lettres de stage, renouvellement et attestations.
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)

DESCRIPTIONS_DIRECTIONS = {
    "DARH": "Direction des Affaires et Ressources Humaines",
    "DCGIS": "Direction du Centre de Gestion de l'Information et des Statistiques",
    "DEPP": "Direction des Études, Planification et Projets",
    "DT": "Direction Technique",
    "DM": "Direction des Marchés",
    "DFC": "Direction Financière et Comptable",
    "DG": "Direction Générale",
}

SIGNATAIRES = {
    "DG": {"titre": "Le Directeur Général", "nom": "Dr. Karimou CHABI SIKA"},
    "DGA": {"titre": "Le Directeur Général Adjoint", "nom": "M. Damipi NOUPOKOU"},
}


def _create_base_document():
    """Crée un document Word de base avec marges et police configurées."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    return doc


def _add_empty_lines(doc, count=1):
    """Ajoute des lignes vides."""
    for _ in range(count):
        doc.add_paragraph('')


def _add_title(doc, text):
    """Ajoute un titre centré, gras et souligné."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.underline = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    return p


def _add_justified_paragraph(doc, text, bold_parts=None):
    """Ajoute un paragraphe justifié. Supporte le texte avec des parties en gras via des marqueurs <b>...</b>."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)

    # Parser les balises <b>...</b>
    import re
    parts = re.split(r'(<b>.*?</b>)', text)
    for part in parts:
        if part.startswith('<b>') and part.endswith('</b>'):
            run = p.add_run(part[3:-4])
            run.bold = True
        else:
            p.add_run(part)

    for run in p.runs:
        run.font.size = Pt(12)
        run.font.name = 'Calibri'

    return p


def _add_signature(doc, signataire_titre, signataire_nom):
    """Ajoute un bloc de signature aligné à droite."""
    _add_empty_lines(doc, 3)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = p1.add_run(signataire_titre)
    run1.font.size = Pt(12)
    run1.font.name = 'Calibri'

    _add_empty_lines(doc, 2)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(signataire_nom)
    run2.bold = True
    run2.font.size = Pt(12)
    run2.font.name = 'Calibri'


def _get_direction_description(direction_code):
    """Retourne la description complète d'une direction."""
    return DESCRIPTIONS_DIRECTIONS.get(direction_code, direction_code)


def generer_lettre_stage_docx(demande, donnees, signataire_id="DG", helpers=None):
    """
    Génère la lettre d'autorisation de stage en Word (.docx).

    Args:
        demande: objet Demande
        donnees: dict avec date_debut, date_fin, direction, service, remunere, montant
        signataire_id: "DG" ou "DGA"
        helpers: dict avec les fonctions utilitaires (calculer_duree_mois, date_format_francais, etc.)

    Returns:
        BytesIO buffer contenant le document Word
    """
    sig = SIGNATAIRES.get(signataire_id, SIGNATAIRES["DG"])

    duree_mois = helpers["calculer_duree_mois"](donnees["date_debut"], donnees["date_fin"])
    date_debut_format = helpers["date_format_francais"](donnees["date_debut"])
    date_fin_format = helpers["date_format_francais"](donnees["date_fin"])
    duree_formatee = helpers["duree_en_lettres_et_chiffres"](duree_mois)

    direction_code = donnees.get("direction", "")
    direction_description = _get_direction_description(direction_code)

    doc = _create_base_document()

    # Espace pour en-tête pré-imprimé
    _add_empty_lines(doc, 10)

    # Titre
    _add_title(doc, "AUTORISATION DE STAGE")
    _add_empty_lines(doc, 1)

    # Texte principal
    _add_justified_paragraph(
        doc,
        f"Monsieur <b>{demande.etudiant_nom.upper()} {demande.etudiant_prenom}</b> est autorisé à effectuer un stage à la CEB "
        f"pour une durée de <b>{duree_formatee}</b> mois allant du <b>{date_debut_format}</b> "
        f"au <b>{date_fin_format}</b>."
    )
    _add_empty_lines(doc, 1)

    _add_justified_paragraph(
        doc,
        f"L'intéressé est mis à la disposition de la Direction <b>{direction_code} ({direction_description})</b>, "
        f"Service <b>{donnees.get('service', '')}</b>."
    )
    _add_empty_lines(doc, 1)

    # Rémunération
    if donnees.get("remunere", False):
        montant = donnees.get("montant", 0)
        montant_formate = f"{montant:,}".replace(",", " ")
        _add_justified_paragraph(
            doc,
            f"Conformément à la décision Nº236/CEB/DG/DARH/SAA/SASR/2020 portant Révision des indemnités "
            f"forfaitaires de Stage en date du 3 septembre 2020, Monsieur <b>{demande.etudiant_nom.upper()} {demande.etudiant_prenom}</b> "
            f"bénéficie au cours de son stage d'une indemnité forfaitaire mensuelle nette de "
            f"<b>{montant_formate} FRANCS CFA</b>."
        )

    # Signature
    _add_signature(doc, sig["titre"], sig["nom"])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generer_lettre_renouvellement_docx(stagiaire, donnees, signataire_id="DG", helpers=None, convention_numero=None):
    """
    Génère la lettre d'autorisation de renouvellement de stage en Word (.docx).

    Args:
        stagiaire: objet Stagiaire
        donnees: dict avec date_debut, date_fin, direction, service, remunere, montant
        signataire_id: "DG" ou "DGA"
        helpers: dict avec fonctions utilitaires
        convention_numero: numéro de convention pour le pied de page

    Returns:
        BytesIO buffer
    """
    sig = SIGNATAIRES.get(signataire_id, SIGNATAIRES["DG"])

    date_debut_str = donnees.get("date_debut") or donnees.get("date_debut_iso", "")
    date_fin_str = donnees.get("date_fin") or donnees.get("date_fin_iso", "")

    duree_mois = helpers["calculer_duree_mois"](date_debut_str, date_fin_str)
    duree_formatee = helpers["duree_en_lettres_et_chiffres"](duree_mois)

    date_debut_format = helpers["date_format_francais"](date_debut_str)
    date_fin_format = helpers["date_format_francais"](date_fin_str)

    # Dates du stage précédent
    date_debut_precedent = helpers["date_format_francais"](
        stagiaire.date_debut.strftime("%d/%m/%Y")
    )
    date_fin_precedent = helpers["date_format_francais"](
        stagiaire.date_fin.strftime("%d/%m/%Y")
    )

    direction_code = donnees.get("direction", "")
    direction_description = _get_direction_description(direction_code)

    doc = _create_base_document()

    # Espace pour en-tête pré-imprimé
    _add_empty_lines(doc, 10)

    # Titre
    _add_title(doc, "AUTORISATION DE RENOUVELLEMENT DE STAGE")
    _add_empty_lines(doc, 1)

    # Texte principal
    _add_justified_paragraph(
        doc,
        f"Dans le cadre du renouvellement de son stage à la CEB, "
        f"Monsieur <b>{stagiaire.nom.upper()} {stagiaire.prenom}</b> est autorisé à effectuer un nouveau stage "
        f"pour une durée de <b>{duree_formatee}</b> mois allant du <b>{date_debut_format}</b> "
        f"au <b>{date_fin_format}</b>."
    )
    _add_empty_lines(doc, 1)

    _add_justified_paragraph(
        doc,
        f"L'intéressé est mis à la disposition de la Direction <b>{direction_code} ({direction_description})</b>, "
        f"Service <b>{donnees.get('service', '')}</b>."
    )
    _add_empty_lines(doc, 1)

    # Mention du stage précédent
    _add_justified_paragraph(
        doc,
        f"Ce renouvellement fait suite au stage précédent effectué du "
        f"<b>{date_debut_precedent}</b> au <b>{date_fin_precedent}</b>."
    )
    _add_empty_lines(doc, 1)

    # Rémunération
    if donnees.get("remunere", False):
        montant = donnees.get("montant", 0)
        montant_formate = f"{montant:,}".replace(",", " ")
        _add_justified_paragraph(
            doc,
            f"Conformément à la décision Nº236/CEB/DG/DARH/SAA/SASR/2020 portant Révision des indemnités "
            f"forfaitaires de Stage en date du 3 septembre 2020, Monsieur <b>{stagiaire.nom.upper()} {stagiaire.prenom}</b> "
            f"bénéficie au cours de son stage d'une indemnité forfaitaire mensuelle nette de "
            f"<b>{montant_formate} FRANCS CFA</b>."
        )

    # Signature
    _add_signature(doc, sig["titre"], sig["nom"])

    # Pied de page
    if convention_numero:
        _add_empty_lines(doc, 2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CONVENTION DE RENOUVELLEMENT TEMPORAIRE - Réf: {convention_numero}")
        run.italic = True
        run.font.size = Pt(8)
        run.font.name = 'Calibri'
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generer_attestation_docx(demande_attestation, signataire_id="DG"):
    """
    Génère l'attestation de stage en Word (.docx),
    avec le même format que la lettre de stage.

    Args:
        demande_attestation: objet DemandeAttestation
        signataire_id: "DG" ou "DGA"

    Returns:
        BytesIO buffer
    """
    from django.utils import timezone

    sig = SIGNATAIRES.get(signataire_id, SIGNATAIRES["DG"])
    stagiaire = demande_attestation.stagiaire

    date_debut = stagiaire.date_debut.strftime("%d/%m/%Y") if stagiaire.date_debut else "__________"
    date_fin = stagiaire.date_fin.strftime("%d/%m/%Y") if stagiaire.date_fin else "__________"

    duree_mois = "____"
    if stagiaire.date_debut and stagiaire.date_fin:
        delta = stagiaire.date_fin - stagiaire.date_debut
        mois = delta.days // 30
        duree_mois = f"{mois:02d}" if mois > 0 else "____"

    direction_code = stagiaire.direction or "__________"
    direction_description = _get_direction_description(direction_code)

    doc = _create_base_document()

    # Espace pour en-tête pré-imprimé
    _add_empty_lines(doc, 10)

    # Titre
    _add_title(doc, "ATTESTATION DE STAGE")
    _add_empty_lines(doc, 1)

    # Texte principal
    _add_justified_paragraph(
        doc,
        f"Nous soussignés, attestons par la présente que <b>{stagiaire.nom.upper()} {stagiaire.prenom}</b>, "
        f"dans le cadre de son perfectionnement en <b>{stagiaire.specialite or '__________'}</b>, "
        f"a effectué un stage de <b>{duree_mois}</b> mois, "
        f"du <b>{date_debut}</b> au <b>{date_fin}</b> "
        f"au sein de la <b>{direction_code} ({direction_description})</b>."
    )
    _add_empty_lines(doc, 1)

    _add_justified_paragraph(
        doc,
        f"Au cours de son stage effectué avec assiduité, "
        f"<b>{stagiaire.nom.upper()} {stagiaire.prenom}</b> s'est montré dévoué "
        f"et a attaché un grand intérêt au travail bien fait."
    )
    _add_empty_lines(doc, 1)

    _add_justified_paragraph(
        doc,
        "En foi de quoi, la présente attestation lui est délivrée "
        "pour servir et valoir ce que de droit."
    )

    # Signature
    _add_signature(doc, sig["titre"], sig["nom"])

    # Pied de page
    _add_empty_lines(doc, 2)
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    from docx.shared import RGBColor
    date_generation = timezone.now().strftime("Généré le %d/%m/%Y à %H:%M")
    run_date = p_footer.add_run(date_generation)
    run_date.font.size = Pt(9)
    run_date.font.name = 'Calibri'
    run_date.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Ajouter un tab et la référence
    run_tab = p_footer.add_run("\t\t\t\t\t\t")
    run_ref = p_footer.add_run(f"Réf: ATT-{demande_attestation.id:06d}")
    run_ref.font.size = Pt(9)
    run_ref.font.name = 'Calibri'
    run_ref.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
